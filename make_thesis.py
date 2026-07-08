#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)

def create_massive_document():
    doc = Document()
    
    # --- TITLE PAGE ---
    title = doc.add_heading('SpectralGraph', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('The Complete Engineering Manual: From Zero to Hero')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n\n\n')
    doc.add_paragraph('Author: Zuni').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Program: BSc Computer Science (4th Semester)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Institution: Abbottabad University of Science & Technology (AUST)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- CHAPTER 1 ---
    doc.add_heading('Chapter 1: The Problem with Cybersecurity Data', level=1)
    doc.add_paragraph('To understand SpectralGraph, you must first understand the problem it solves. During a "Red Team" engagement (where ethical hackers simulate cyber attacks to find weaknesses), operators use tools like Nmap to scan networks. These tools find IP addresses, open ports, and operating systems.')
    doc.add_paragraph('Traditionally, hackers copy and paste these results into Excel spreadsheets. This is highly inefficient. If a target opens a new port the next day, updating the spreadsheet overwrites the old data, destroying the historical timeline. Furthermore, if two machines have the same vulnerability, the operator must type the vulnerability details twice. SpectralGraph solves this by introducing a NoSQL database architecture to automate and map this data.')
    
    # --- CHAPTER 2 ---
    doc.add_heading('Chapter 2: Why MongoDB? (NoSQL vs SQL)', level=1)
    doc.add_paragraph('In a traditional SQL database, data is forced into rigid tables with strict columns. However, network scan data is hierarchical and messy. One machine might have 2 open ports, while another has 500. MongoDB, a NoSQL Document database, stores data in flexible JSON-like documents.')
    doc.add_paragraph('SpectralGraph uses a "Hybrid Schema":')
    doc.add_paragraph('1. Embedding: We embed the "open ports" directly inside the target document. This means when we query a target, we instantly get its entire attack surface without needing complex JOIN operations.', style='List Number')
    doc.add_paragraph('2. Referencing: We store "vulnerabilities" in a separate collection, linking them to targets via a target_id. This normalizes the data, so a specific CVE (Common Vulnerability) is only written once but can be linked to thousands of machines.', style='List Number')
    doc.add_page_break()

    # --- CHAPTER 3 ---
    doc.add_heading('Chapter 3: Module 1 - Automated Ingestion & Idempotency', level=1)
    doc.add_paragraph('The core of SpectralGraph is its Python ingestion engine. Instead of manual data entry, the python-nmap library scans the network and passes the raw data to PyMongo.')
    doc.add_paragraph('Idempotency is a computer science concept meaning an operation can be applied multiple times without changing the result beyond the initial application. In SpectralGraph, if we scan the same IP address 50 times, we do not want 50 duplicate target records. We achieve idempotency using MongoDB\'s "upsert=True" command. If the IP exists, it updates the ports. If it does not exist, it creates a new document.')
    doc.add_paragraph('Core Ingestion Logic snippet:')
    add_code(doc, 'db.targets.update_one(\n    {"ip_address": ip},\n    {"$set": target_document},\n    upsert=True\n)')
    
    # --- CHAPTER 4 ---
    doc.add_heading('Chapter 4: Module 2 - Time-Series Audit Logging', level=1)
    doc.add_paragraph('While the targets collection tracks the "Current State" of the network, security audits require historical tracking. To achieve this, every time the Python script runs, it also performs an insert_one() operation into the scan_history collection.')
    doc.add_paragraph('This creates an immutable (unchangeable) time-series log. An analyst can look at the database and say, "Port 22 was closed on Tuesday, but opened on Thursday," proving a lateral movement or network change.')
    add_code(doc, 'history_document = {\n    "target_id": target_record["_id"],\n    "scanned_at": datetime.now(timezone.utc),\n    "scan_parameters": "-sV -O",\n    "status_at_time": state,\n    "ports_found": ports_list\n}\ndb.scan_history.insert_one(history_document)')
    doc.add_page_break()

    # --- CHAPTER 5 ---
    doc.add_heading('Chapter 5: Module 3 - Role-Based Access Control (RBAC)', level=1)
    doc.add_paragraph('A database is useless if anyone can modify it. SpectralGraph implements application-layer security. It features a users collection where passwords are encrypted using SHA-256 hashing (via Python\'s hashlib).')
    doc.add_paragraph('When a user attempts to log a new vulnerability, the add_vuln.py script intercepts the request. It verifies the password hash, and then checks the user\'s role. If the user is an "Analyst", the database transaction is rejected. Only a "Lead_Operator" is authorized to write to the vulnerabilities collection.')
    
    # --- CHAPTER 6 ---
    doc.add_heading('Chapter 6: Module 4 - Advanced Aggregation Pipelines', level=1)
    doc.add_paragraph('To generate executive dashboards, SpectralGraph bypasses standard read queries and utilizes MongoDB Aggregation Pipelines. This 4-stage pipeline is the most advanced concept in the project:')
    doc.add_paragraph('Stage 1 ($lookup): Acts as a NoSQL LEFT JOIN, fetching all vulnerabilities linked to a target\'s _id.', style='List Bullet')
    doc.add_paragraph('Stage 2 ($unwind): Deconstructs the vulnerability array so each CVE becomes its own readable row.', style='List Bullet')
    doc.add_paragraph('Stage 3 ($group): Dynamically groups the data by Operating System AND Severity, counting the exact number of critical flaws per platform.', style='List Bullet')
    doc.add_paragraph('Stage 4 ($sort): Orders the output logically for the terminal dashboard.', style='List Bullet')
    
    # --- CONCLUSION ---
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('SpectralGraph successfully bridges the gap between offensive cybersecurity operations and modern data architecture. By migrating from flat spreadsheets to a fully automated, relational NoSQL backend, red teams can track targets faster, maintain historical integrity, and query massive datasets instantly.')

    # Save
    doc.save('SpectralGraph_Complete_Manual.docx')
    print("[+] Massive Manual successfully saved as 'SpectralGraph_Complete_Manual.docx'")

if __name__ == '__main__':
    create_massive_document()
