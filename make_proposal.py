#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt

def create_proposal():
    doc = Document()
    
    # Title
    title = doc.add_heading('Project Proposal', 0)
    
    # Header Info
    doc.add_paragraph().add_run('Project Title: SpectralGraph: A NoSQL-Backed Red Team Asset & Vulnerability Tracker').bold = True
    doc.add_paragraph().add_run('Student Name: Zuni').bold = True
    doc.add_paragraph().add_run('Program: BSc Computer Science').bold = True
    doc.add_paragraph().add_run('Semester: 4th Semester').bold = True
    doc.add_paragraph().add_run('Institution: Abbottabad University of Science & Technology (AUST)').bold = True
    
    # Section 1
    doc.add_heading('1. Introduction and Problem Statement', level=1)
    doc.add_paragraph('During penetration testing and network administration, security professionals generate massive amounts of unstructured data from tools like Nmap. Traditionally, this data is tracked manually using static spreadsheets. This manual approach creates several critical structural problems:')
    doc.add_paragraph('Lack of Historical Context: Overwriting old scan data destroys the audit trail, making it impossible to see when a specific port was opened or closed.', style='List Bullet')
    doc.add_paragraph('Data Duplication: Flat file structures cannot efficiently map complex, many-to-many relationships.', style='List Bullet')
    doc.add_paragraph('Security Risks: Spreadsheets lack application-layer access control, allowing any user to modify critical vulnerability findings.', style='List Bullet')
    
    # Section 2
    doc.add_heading('2. Proposed Solution', level=1)
    doc.add_paragraph('SpectralGraph is an automated, NoSQL-backed database application designed to resolve the inefficiencies of manual tracking. Built on MongoDB, this system will automatically ingest raw, hierarchical network scan data, track infrastructure changes over time using time-series logging, and securely map identified Common Vulnerabilities and Exposures (CVEs) to specific targets using relational references.')
    
    # Section 3
    doc.add_heading('3. Technology Stack', level=1)
    doc.add_paragraph('Database Engine: MongoDB (NoSQL)', style='List Bullet')
    doc.add_paragraph('Backend Logic: Python 3', style='List Bullet')
    doc.add_paragraph('Integration Libraries: PyMongo, python-nmap', style='List Bullet')
    doc.add_paragraph('Security: hashlib (SHA-256)', style='List Bullet')
    doc.add_paragraph('Environment: Kali Linux', style='List Bullet')
    
    # Section 4
    doc.add_heading('4. Database Architecture & Schema Design', level=1)
    doc.add_paragraph('Unlike rigid relational tables, SpectralGraph utilizes a hybrid NoSQL schema to optimize both read speeds and data normalization. The database will consist of four primary collections:')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Collection'
    hdr_cells[1].text = 'Schema Paradigm'
    hdr_cells[2].text = 'Purpose'
    
    row1 = table.add_row().cells
    row1[0].text = 'targets'
    row1[1].text = 'Embedding'
    row1[2].text = 'Stores the "Current State" of a host. Open ports are embedded directly for rapid read access.'
    
    row2 = table.add_row().cells
    row2[0].text = 'scan_history'
    row2[1].text = 'Referencing'
    row2[2].text = 'Acts as an immutable time-series audit log.'
    
    row3 = table.add_row().cells
    row3[0].text = 'vulnerabilities'
    row3[1].text = 'Referencing'
    row3[2].text = 'Normalizes data by linking specific CVEs to the target ID.'
    
    row4 = table.add_row().cells
    row4[0].text = 'users'
    row4[1].text = 'Standalone'
    row4[2].text = 'Stores authenticated operator credentials and roles.'

    # Section 5
    doc.add_heading('5. Core System Modules', level=1)
    doc.add_paragraph('Automated Data Ingestion (Idempotency): Uses PyMongo upsert logic to automatically update records.', style='List Number')
    doc.add_paragraph('Role-Based Access Control (RBAC): Application-layer constraints restrict database write operations to Lead Operators.', style='List Number')
    doc.add_paragraph('Advanced Analytical Pipelines: Executes multi-stage NoSQL aggregation pipelines to perform complex data analysis.', style='List Number')

    # Section 6
    doc.add_heading('6. Expected Outcomes and Deliverables', level=1)
    doc.add_paragraph('Upon completion, SpectralGraph will demonstrate the practical application of NoSQL document stores. Deliverables include:', style='List Bullet')
    doc.add_paragraph('The fully initialized MongoDB database instance.', style='List Bullet')
    doc.add_paragraph('The Python-based ingestion and analytics engine.', style='List Bullet')
    doc.add_paragraph('A comprehensive project report and presentation.', style='List Bullet')

    # Save it
    doc.save('SpectralGraph_Project_Proposal.docx')
    print("[+] Proposal successfully saved as 'SpectralGraph_Project_Proposal.docx'")

if __name__ == '__main__':
    create_proposal()
