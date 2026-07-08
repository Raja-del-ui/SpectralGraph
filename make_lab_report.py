#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_code_block(doc, text):
    # Safely creates a terminal-looking block without relying on missing styles
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    p.paragraph_format.left_indent = Inches(0.5)

def create_lab_report():
    doc = Document()
    
    # Title and Header
    title = doc.add_heading('Database Systems - Laboratory Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph().add_run('Project: SpectralGraph - Red Team Asset Tracker').bold = True
    doc.add_paragraph().add_run('Student Name: Zuni').bold = True
    doc.add_paragraph().add_run('Program: BSc Computer Science (4th Semester)').bold = True
    doc.add_paragraph().add_run('Institution: Abbottabad University of Science & Technology (AUST)').bold = True
    doc.add_paragraph('_' * 50)
    
    # 1. Objective
    doc.add_heading('1. Lab Objective', level=1)
    doc.add_paragraph('The objective of this laboratory exercise was to implement a NoSQL database architecture using MongoDB to track network assets and vulnerabilities. The lab required demonstrating automated data ingestion via Python, hybrid schema design (embedding and referencing), Role-Based Access Control (RBAC), and advanced multi-stage aggregation pipelines.')
    
    # 2. Environment Setup
    doc.add_heading('2. Environment & Tools', level=1)
    doc.add_paragraph('Operating System: Kali Linux', style='List Bullet')
    doc.add_paragraph('Database Engine: MongoDB Server 7.0', style='List Bullet')
    doc.add_paragraph('Programming Language: Python 3.11+', style='List Bullet')
    doc.add_paragraph('Libraries Used: PyMongo, python-nmap, hashlib', style='List Bullet')
    
    # 3. Implementation Phases
    doc.add_heading('3. Implementation & Results', level=1)
    
    # Phase A: Ingestion
    doc.add_heading('Phase A: Automated Data Ingestion (Idempotency)', level=2)
    doc.add_paragraph('A Python script was developed to execute an Nmap service and OS detection scan. The results were parsed and inserted into the MongoDB database. To prevent duplicate entries during subsequent scans, the PyMongo update_one() function was utilized with upsert=True.')
    
    doc.add_paragraph('Terminal Execution Output:').bold = True
    add_code_block(doc, '┌──(attacker㉿kali)-[~/database-project]\n└─$ sudo python3 nmap_to_mongo.py 10.0.2.15\n[*] Starting intense Nmap scan on: 10.0.2.15...\n[+] Sync\'d 10.0.2.15 to targets AND logged to scan_history.')
    
    # Phase B: RBAC
    doc.add_heading('Phase B: Role-Based Access Control (RBAC)', level=2)
    doc.add_paragraph('To secure the vulnerability data, a users collection was created storing SHA-256 hashed passwords. Application-layer logic was implemented to block users with the "Analyst" role from inserting CVEs, while permitting the "Lead_Operator".')
    
    doc.add_paragraph('Terminal Execution Output (Admin Login):').bold = True
    add_code_block(doc, '└─$ python3 add_vuln.py 10.0.2.15 "CVE-2024-3094" "XZ Utils Backdoor" "Critical"\n--- SECURE LOGIN REQUIRED ---\nUsername: zuni_admin\nPassword: \n[+] Authentication and Authorization successful.\n[+] Vulnerability \'XZ Utils Backdoor\' successfully linked to 10.0.2.15 by zuni_admin')
    
    # Phase C: Aggregation
    doc.add_heading('Phase C: Advanced Data Aggregation', level=2)
    doc.add_paragraph('A 4-stage MongoDB aggregation pipeline was executed to act as a NoSQL LEFT JOIN, combining the targets and vulnerabilities collections. The pipeline utilized $lookup, $unwind, $group, and $sort to generate a statistical dashboard.')
    
    doc.add_paragraph('Terminal Execution Output:').bold = True
    add_code_block(doc, '============================================================\n 📊 NETWORK SECURITY ANALYTICS: VULNERABILITIES BY OS\n============================================================\n\n[+] Operating System: Unknown\n    -> Critical Severity: 2 CVE(s)')

    # 4. Conclusion
    doc.add_heading('4. Conclusion', level=1)
    doc.add_paragraph('The laboratory exercise successfully demonstrated how NoSQL databases can handle hierarchical network data more efficiently than flat files. The implementation of time-series logging via the scan_history collection and relational linking via the target_id field proved that MongoDB can support robust, secure, and highly relational application backends.')

    # Save it
    doc.save('SpectralGraph_Lab_Report.docx')
    print("[+] Lab Report successfully saved as 'SpectralGraph_Lab_Report.docx'")

if __name__ == '__main__':
    create_lab_report()
